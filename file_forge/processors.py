"""
File processing engine with sorting, conversion, and analysis capabilities.
"""

import hashlib
import mimetypes
from pathlib import Path
from typing import Dict, Any, List
from PIL import Image
import logging
from datetime import datetime

# Optional dependencies for specialized file processing
try:
    import moviepy.editor as mp
except ImportError:
    mp = None

try:
    from pydub import AudioSegment
except ImportError:
    AudioSegment = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import docx
except ImportError:
    docx = None

from .models import File, SortingRule
from .specialized_processors import specialized_processor

logger = logging.getLogger(__name__)


class FileProcessor:
    """Main file processing engine."""

    def __init__(self):
        """Initialize processor with supported formats."""
        self.supported_formats = {
            'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp'],
            'video': ['.mp4', '.avi', '.mov', '.wmv', '.flv', '.mkv'],
            'audio': ['.mp3', '.wav', '.flac', '.aac', '.ogg', '.wma'],
            'document': ['.pdf', '.docx', '.doc', '.txt', '.rtf', '.odt'],
            'archive': ['.zip', '.rar', '.7z', '.tar', '.gz'],
            # Specialized formats
            'cad': ['.dwg', '.dxf', '.step', '.stp', '.iges'],
            'ebook': ['.epub', '.mobi', '.fb2', '.html', '.xhtml', '.chm'],
            'medical': ['.dcm', '.dicom', '.hl7', '.xhl7', '.pdfa'],
            'geospatial': ['.shp', '.geojson', '.svg', '.kml', '.kmz', '.gpx', '.gml'],
            'disk_image': ['.dmg', '.iso', '.img', '.vhd', '.vmdk'],
            'email': ['.eml', '.msg', '.pst', '.mbox']
        }
        self._specialized_types = {'cad', 'ebook', 'medical', 'geospatial', 'disk_image', 'email'}

    def process_file(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process a file and extract metadata."""
        logger.info(f"Processing file: {filename}")

        file_path_obj = Path(file_path)
        if not file_path_obj.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Basic file info
        stat = file_path_obj.stat()
        file_size = stat.st_size
        modified_time = datetime.fromtimestamp(stat.st_mtime)

        # MIME type and file type detection
        mime_type, _ = mimetypes.guess_type(filename)
        file_type = self._determine_file_type(filename)

        # Calculate file hash
        file_hash = self._calculate_file_hash(file_path)

        # Check if specialized processor can handle this file
        if specialized_processor.can_process(filename):
            try:
                specialized_result = specialized_processor.process(file_path, filename)
                # Merge specialized result with basic info
                specialized_result['file_size'] = file_size
                specialized_result['file_hash'] = file_hash
                specialized_result['modified_time'] = modified_time
                return specialized_result
            except Exception as e:
                logger.warning(f"Specialized processing failed for {filename}: {e}")

        # Extract metadata based on file type
        metadata = self._extract_metadata(file_path, filename, file_type)

        # Extract text content if possible
        content = self._extract_content(file_path, filename, file_type)

        return {
            'filename': filename,
            'file_path': file_path,
            'file_type': file_type,
            'mime_type': mime_type,
            'file_size': file_size,
            'file_hash': file_hash,
            'modified_time': modified_time,
            'metadata': metadata,
            'content': content
        }

    def _determine_file_type(self, filename: str) -> str:
        """Determine file type from extension."""
        ext = Path(filename).suffix.lower()

        for category, extensions in self.supported_formats.items():
            if ext in extensions:
                return category

        return 'unknown'

    def _calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA-256 hash of file."""
        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()

    def _extract_metadata(self, file_path: str, filename: str, file_type: str) -> Dict[str, Any]:
        """Extract metadata based on file type."""
        metadata = {}

        try:
            if file_type == 'image':
                metadata.update(self._extract_image_metadata(file_path))
            elif file_type == 'video':
                metadata.update(self._extract_video_metadata(file_path))
            elif file_type == 'audio':
                metadata.update(self._extract_audio_metadata(file_path))
            elif file_type == 'document':
                metadata.update(self._extract_document_metadata(file_path, filename))
        except Exception as e:
            logger.warning(f"Failed to extract metadata for {filename}: {e}")

        return metadata

    def _extract_image_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract image metadata."""
        with Image.open(file_path) as img:
            has_alpha = img.mode in ('RGBA', 'LA')
            if img.mode == 'P' and 'transparency' in img.info:
                has_alpha = True
            
            return {
                'width': img.width,
                'height': img.height,
                'format': img.format,
                'mode': img.mode,
                'has_alpha': has_alpha,
            }

    def _extract_video_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract video metadata."""
        try:
            clip = mp.VideoFileClip(file_path)
            metadata = {
                'duration': clip.duration,
                'width': clip.w,
                'height': clip.h,
                'fps': clip.fps,
            }
            clip.close()
            return metadata
        except Exception:
            return {}

    def _extract_audio_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract audio metadata."""
        try:
            audio = AudioSegment.from_file(file_path)
            return {
                'duration': len(audio) / 1000,  # Convert to seconds
                'channels': audio.channels,
                'sample_width': audio.sample_width,
                'frame_rate': audio.frame_rate,
                'max_amplitude': audio.max,
            }
        except Exception:
            return {}

    def _extract_document_metadata(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Extract document metadata."""
        metadata = {}
        ext = Path(filename).suffix.lower()

        try:
            if ext == '.pdf':
                with open(file_path, 'rb') as f:
                    pdf = PyPDF2.PdfReader(f)
                    metadata = {
                        'pages': len(pdf.pages),
                        'title': pdf.metadata.get('/Title', ''),
                        'author': pdf.metadata.get('/Author', ''),
                        'subject': pdf.metadata.get('/Subject', ''),
                        'creator': pdf.metadata.get('/Creator', ''),
                    }
            elif ext == '.docx':
                doc = docx.Document(file_path)
                metadata = {
                    'paragraphs': len(doc.paragraphs),
                    'word_count': sum(len(p.text.split()) for p in doc.paragraphs),
                }
        except Exception:
            pass

        return metadata

    def _extract_content(self, file_path: str, filename: str, file_type: str) -> str:
        """Extract text content from file if possible."""
        content = ""
        ext = Path(filename).suffix.lower()

        try:
            if ext == '.txt':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
            elif ext == '.pdf':
                with open(file_path, 'rb') as f:
                    pdf = PyPDF2.PdfReader(f)
                    for page in pdf.pages:
                        content += page.extract_text() + "\n"
            elif ext == '.docx':
                doc = docx.Document(file_path)
                content = "\n".join([p.text for p in doc.paragraphs])
        except Exception as e:
            logger.warning(f"Failed to extract content from {filename}: {e}")

        return content


class SortingEngine:
    """Advanced sorting engine with customizable rules."""

    def __init__(self):
        """Initialize sorting engine."""
        self.file_processor = FileProcessor()

    def apply_sorting_rules(self, files: List[File], rules: List[SortingRule]) -> List[File]:
        """Apply sorting rules to a list of files."""
        if not rules:
            return files

        sorted_files = files.copy()

        for rule in rules:
            if not rule.is_active:
                continue

            try:
                sorted_files = self._apply_single_rule(sorted_files, rule)
            except Exception as e:
    def _apply_single_rule(self, files: List[File], rule: SortingRule) -> List[File]:
        """Apply a single sorting rule."""
        conditions = rule.conditions or {}
        actions = rule.actions or {}

        # Filter files based on conditions
        matching_files = []
        non_matching_files = []
        for file in files:
            if self._matches_conditions(file, conditions):
                matching_files.append(file)
            else:
                non_matching_files.append(file)

        # Apply sorting action
        sort_by = actions.get('sort_by')
        sort_order = actions.get('sort_order', 'asc')

        if sort_by:
            reverse = sort_order == 'desc'
            matching_files.sort(key=lambda f: self._get_sort_key(f, sort_by), reverse=reverse)

        # Return matching files first, then non-matching
        return matching_files + non_matching_files
            reverse = sort_order == 'desc'
            filtered_files.sort(key=lambda f: self._get_sort_key(f, sort_by), reverse=reverse)

        return filtered_files

    def _matches_conditions(self, file: File, conditions: Dict[str, Any]) -> bool:
        """Check if file matches rule conditions."""
        for key, value in conditions.items():
            if key == 'file_type' and file.file_type != value:
                return False
            elif key == 'file_size_min' and file.file_size < value:
                return False
            elif key == 'file_size_max' and file.file_size > value:
                return False
            elif key == 'created_after' and file.created_at < value:
                return False
            elif key == 'created_before' and file.created_at > value:
                return False
            elif key == 'has_tag':
                if not file.tags or value not in file.tags:
                    return False

        return True

    def _get_sort_key(self, file: File, sort_by: str):
        """Get sorting key for file."""
        if sort_by == 'filename':
            return file.filename.lower()
        elif sort_by == 'file_size':
            return file.file_size
        elif sort_by == 'file_type':
            return file.file_type
        elif sort_by == 'created_at':
            return file.created_at
        elif sort_by == 'modified_time':
            return file.modified_time if hasattr(file, 'modified_time') else file.created_at
        else:
            return file.filename.lower()  # Default fallback


# Global instances
file_processor = FileProcessor()
sorting_engine = SortingEngine()